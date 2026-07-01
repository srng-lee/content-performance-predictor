"""
마크다운(.md) 리포트를 워드(.docx)로 변환하는 범용 스크립트.

사용법:
    python tools/md_to_docx.py <입력.md> <출력.docx>

지원 문법: #~###### 제목, **굵게**, - 불릿, 1. 번호 목록, > 인용구,
표(| ... |), --- 구분선. 그 외 문법은 없는 것으로 취급하고 일반 문단으로 처리한다.
"""

import re
import sys

from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8", errors="replace")


def add_runs_with_bold(paragraph, text):
    """**bold** 구문을 파싱해 굵게 표시된 run으로 나눠 추가한다."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def is_table_row(line):
    return line.startswith("|") and line.endswith("|")


def is_table_separator(line):
    return bool(re.match(r"^\|?[\s:\-|]+\|?$", line)) and "-" in line


def parse_table_row(line):
    return [c.strip() for c in line.strip("|").split("|")]


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)

    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        if is_table_row(stripped) and i + 1 < n and is_table_separator(lines[i + 1].strip()):
            header = parse_table_row(stripped)
            i += 2
            rows = []
            while i < n and is_table_row(lines[i].strip()):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, text in enumerate(header):
                run = table.rows[0].cells[c].paragraphs[0].add_run(text)
                run.bold = True
            for row_vals in rows:
                cells = table.add_row().cells
                for c, text in enumerate(row_vals):
                    if c < len(cells):
                        add_runs_with_bold(cells[c].paragraphs[0], text)
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs_with_bold(p, stripped.lstrip(">").strip())
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, bullet_match.group(1))
            i += 1
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, numbered_match.group(2))
            i += 1
            continue

        indented_bullet = re.match(r"^\s{2,}-\s+(.*)$", lines[i])
        if indented_bullet:
            p = doc.add_paragraph(style="List Bullet 2")
            add_runs_with_bold(p, indented_bullet.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"saved: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("사용법: python tools/md_to_docx.py <입력.md> <출력.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
